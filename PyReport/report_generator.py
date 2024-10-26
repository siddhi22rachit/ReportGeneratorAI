
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import requests
from datetime import datetime
import os

class AIReportGenerator:
    def __init__(self, api_key):
        self.api_key = api_key
        self.doc = Document()
        self.team_members = []
        self.setup_document()

    def setup_document(self):
        """Setup document margins and default styles"""
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            
            header = section.header
            header_para = header.paragraphs[0]
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if os.path.exists('thakur_logo.jpg'):
                run = header_para.add_run()
                run.add_picture('thakur_logo.jpg', width=Inches(2))

    def collect_team_members(self):
        """Collect team member details"""
        num_members = int(input("\nEnter number of team members: "))
        for i in range(num_members):
            print(f"\nEnter details for team member {i+1}:")
            member = {
                'id': input("Student ID: ").strip(),
                'name': input("Name: ").strip(),
                'subject': input("Subject: ").strip()
            }
            self.team_members.append(member)

    def add_page_number(self, paragraph):
        """Add page numbers to the footer"""
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar)
        
        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"
        run._r.append(instrText)
        
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar)

    def add_title_page(self, topic, guide_name):
        """Add title page with project and team details"""
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run("THAKUR COLLEGE OF ENGINEERING AND TECHNOLOGY")
        run.bold = True
        run.font.size = Pt(16)
        
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run("\nDEPARTMENT OF COMPUTER ENGINEERING")
        run.bold = True
        run.font.size = Pt(14)
        
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(f"\n\n{topic.upper()}")
        run.bold = True
        run.font.size = Pt(14)
        
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run("\nA PROJECT REPORT")
        run.font.size = Pt(12)
        
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run("\n\nPrepared by:\n")
        
        for member in self.team_members:
            para.add_run(f"\n{member['name']} ({member['id']})\n{member['subject']}\n")
        
        para.add_run(f"\n\nUnder the guidance of:\n{guide_name}")
        
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run(f"\n\nACADEMIC YEAR {datetime.now().year}-{datetime.now().year + 1}")
        
        self.doc.add_page_break()

    def add_certificate(self, topic, guide_name):
        """Add certificate page"""
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run("CERTIFICATE\n\n").bold = True
        
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        team_members_text = ", ".join([member['name'] for member in self.team_members])
        
        certificate_text = (
            f"This is to certify that the project report entitled \"{topic}\" submitted by "
            f"{team_members_text} to Thakur College of Engineering and Technology, Mumbai, "
            f"is a record of authentic work carried out under my supervision and guidance. "
            f"This work has not been submitted elsewhere for any other degree."
        )
        para.add_run(certificate_text)
        
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.add_run(f"\nDate: {datetime.now().strftime('%d/%m/%Y')}")
        
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        para.add_run(f"\n\n\n{guide_name}\nProject Guide")
        
        self.doc.add_page_break()

    def add_declaration(self, topic):
        """Add declaration page"""
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run("DECLARATION\n\n").bold = True
        
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        team_members_text = ", ".join([member['name'] for member in self.team_members])
        
        declaration_text = (
            f"We hereby declare that the project report entitled \"{topic}\", submitted by us "
            f"to Thakur College of Engineering and Technology, Mumbai, is a result of our own "
            f"work and effort. We have not copied any part from other people's work or from "
            f"any other sources except where due reference or acknowledgment is made explicitly "
            f"in the text. We are aware that any violation of the above will be considered as "
            f"plagiarism and may result in cancellation of the degree."
        )
        para.add_run(declaration_text)
        
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.add_run(f"\nDate: {datetime.now().strftime('%d/%m/%Y')}")
        
        for member in self.team_members:
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            para.add_run(f"\n\n\n{member['name']}")
        
        self.doc.add_page_break()

    def add_acknowledgement(self, guide_name):
        """Add acknowledgement page"""
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run("ACKNOWLEDGEMENT\n\n").bold = True
        
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        acknowledgement_text = (
            f"We would like to express our sincere gratitude to our project guide {guide_name} "
            f"for their invaluable guidance, constant encouragement, and generous help throughout "
            f"the development of this project. Their expertise and insights have been instrumental "
            f"in shaping this work.\n\n"
            f"We are grateful to the Head of Department and all faculty members of the Computer "
            f"Engineering Department for providing the necessary infrastructure and resources. "
            f"Their support has been crucial for the successful completion of this project.\n\n"
            f"We would also like to thank our family and friends for their continuous support and "
            f"encouragement throughout our academic journey."
        )
        para.add_run(acknowledgement_text)
        
        for member in self.team_members:
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            para.add_run(f"\n{member['name']}")
        
        self.doc.add_page_break()

    def add_abstract(self, topic):
        """Generate and add abstract page"""
        abstract_content = self.generate_ai_content(f"Write a technical abstract for {topic} in 250 words")[0]['content']
        
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run("ABSTRACT\n\n").bold = True
        
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.add_run(abstract_content)
        
        self.doc.add_page_break()

    def add_table_of_contents(self):
        """Add table of contents page"""
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run("TABLE OF CONTENTS\n\n").bold = True
        
        table = self.doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        header_cells = table.rows[0].cells
        header_cells[0].text = "Content"
        header_cells[1].text = "Page No."
        
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        chapters = [
            ("1", "Introduction", "1"),
            ("2", "Literature Review", "5"),
            ("3", "System Analysis", "10"),
            ("4", "System Design", "15"),
            ("5", "Implementation", "20"),
            ("6", "Testing and Results", "25"),
            ("7", "Conclusion and Future Scope", "30"),
            ("", "References", "35")
        ]
        
        for chapter in chapters:
            row_cells = table.add_row().cells
            if chapter[0]:
                row_cells[0].text = f"Chapter {chapter[0]}: {chapter[1]}"
            else:
                row_cells[0].text = chapter[1]
            row_cells[1].text = chapter[2]
        
        self.doc.add_page_break()

    def generate_ai_content(self, prompt):
        """Generate content using Google's Gemini API"""
        GEMINI_API_ENDPOINT = "https://generativelanguage.googleapis.com/v1beta/models/gemini-pro:generateContent"
        
        headers = {
            "Content-Type": "application/json",
            "x-goog-api-key": self.api_key
        }
        
        data = {
            "contents": [{
                "parts": [{
                    "text": prompt
                }]
            }]
        }
        
        try:
            response = requests.post(
                GEMINI_API_ENDPOINT,
                headers=headers,
                json=data,
                timeout=30
            )
            response.raise_for_status()
            
            result = response.json()
            if "candidates" not in result or not result["candidates"]:
                raise ValueError("No content generated from API")
                
            text = result["candidates"][0]["content"]["parts"][0]["text"]
            return [{"title": "", "content": text}]
            
        except requests.Timeout:
            print("Request timed out. Please try again.")
        except requests.RequestException as e:
            print(f"Network error: {str(e)}")
        except (KeyError, ValueError) as e:
            print(f"Error parsing API response: {str(e)}")
        except Exception as e:
            print(f"Unexpected error: {str(e)}")
        
        return [{"title": "", "content": "Error generating content. Please try again."}]

    def add_content_sections(self, topic):
        """Add content sections with proper formatting"""
        sections = [
            {
                "title": "Introduction",
                "prompt": f"Write a 100 words report based introduction for {topic}  avoid using bold titel just give para "
            },
            {
                "title": "Literature Review",
                "prompt": f"Write a comprehensive Theoretical Background  for {topic} discussing existing systems, technologies, and research papers. and try to avoid using bold titel just give paraand if available add images also "
            },
            {
                "title": "System Analysis",
                "prompt": f"Write a detailed system analysis for {topic} including  analysis of {topic},  study on this {topic}  .avoid using bold titel just give para"
            },
            {
                "title": "System Design",
                "prompt": f"Write a comprehensive system design section for {topic} including {topic} design, already available system, and interface design. avoid using bold titel just give para"
            },
            {
                "title": "Implementation",
                "prompt": f"Write a detailed implementation section for {topic} including technologies used, development process, and algorithms.avoid using bold titel just give para"
            },
            {
                "title": "Conclusion and Future Scope",
                "prompt": f"Write a conclusion and future scope section for {topic} summarizing achievements and future enhancements."
            }, {
                "title": "Referance",
                "prompt": f"Write a referance in IEEE format for {topic} in points."
            }
        ]
        
        for i, section in enumerate(sections, 1):
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(f"Chapter {i}: {section['title'].upper()}")
            run.bold = True
            run.font.size = Pt(14)
            
            content = self.generate_ai_content(section['prompt'])[0]['content']
            
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.add_run(content)
            
            self.doc.add_page_break()

    def add_references(self):
        """Add references page"""
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run("REFERENCES\n\n").bold = True
        
        references = [
            "IEEE Xplore Digital Library - https://ieeexplore.ieee.org",
            "ACM Digital Library - https://dl.acm.org",
            "Google Scholar - https://scholar.google.com",
            "Research Gate - https://www.researchgate.net"
        ]
        
        for ref in references:
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.add_run(ref)

    def save_document(self, filename):
        """Save document with proper error handling"""
        try:
            filepath = os.path.abspath(filename)
            directory = os.path.dirname(filepath)
            
            if not os.path.exists(directory):
                os.makedirs(directory)
                
            if os.path.exists(filepath):
                backup_file = f"{filepath}.backup"
                os.rename(filepath, backup_file)
                
            self.doc.save(filepath)
            return True
        except Exception as e:
            print(f"Error saving document: {str(e)}")
            return False

    def generate_report(self, topic, guide_name):
        """Generate complete report"""
        try:
            print("\nGenerating report sections...")
            
            self.collect_team_members()
            
            self.add_title_page(topic, guide_name)
            print("✓ Title page added")
            
            self.add_certificate(topic, guide_name)
            print("✓ Certificate added")
            
            self.add_declaration(topic)
            print("✓ Declaration added")
            
            self.add_acknowledgement(guide_name)
            print("✓ Acknowledgement added")
            
            self.add_abstract(topic)
            print("✓ Abstract added")
            
            self.add_table_of_contents()
            print("✓ Table of contents added")
            
            self.add_content_sections(topic)
            print("✓ Main content sections added")
            
            self.add_references()
            print("✓ References added")

            # Add page numbers
            for section in self.doc.sections:
                footer = section.footer
                para = footer.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self.add_page_number(para)
            print("✓ Page numbers added")
            
            # Save document
            filename = f"{topic.replace(' ', '_').lower()}_report.docx"
            if self.save_document(filename):
                print(f"\nSuccess! Report saved as: {filename}")
                return filename
            else:
                print("\nError saving the report.")
                return None
            
        except Exception as e:
            print(f"\nError generating report: {str(e)}")
            return None


def validate_input(topic, guide_name):
    """Validate user input"""
    if not topic or not guide_name:
        return False, "All fields are required"
    
    if len(topic) < 5:
        return False, "Topic should be at least 5 characters long"
    
    if len(guide_name) < 3:
        return False, "Guide name should be at least 3 characters long"
    
    return True, ""


def print_banner():
    """Print application banner"""
    print("\n" + "="*50)
    print("         AI-Powered Project Report Generator")
    print("="*50)
    print("\nThis tool will generate a complete project report using AI")
    print("Make sure you have:")
    print("1. A valid Google Gemini API key")
    print("2. The college logo image (thakur_logo.jpg) in the same directory")
    print("3. Active internet connection")
    print("="*50 + "\n")


def main():
    try:
        print_banner()
        
        # Get API key from environment variable or user input
        api_key = os.getenv('GEMINI_API_KEY')
        if not api_key:
            api_key = input("Enter your Gemini API key: ").strip()
            if not api_key:
                print("Error: API key is required")
                return
        
        # Get user input
        print("\nPlease provide the following details:")
        topic = input("Project Topic: ").strip()
        guide_name = input("Guide Name: ").strip()
        
        # Validate input
        is_valid, error_message = validate_input(topic, guide_name)
        if not is_valid:
            print(f"\nError: {error_message}")
            return
        
        # Create generator instance and generate report
        generator = AIReportGenerator(api_key)
        filename = generator.generate_report(topic, guide_name)
        
        if filename:
            print("\nReport generation completed successfully!")
            print(f"You can find your report at: {os.path.abspath(filename)}")
        else:
            print("\nFailed to generate report. Please check the error messages above.")
            
    except KeyboardInterrupt:
        print("\n\nOperation cancelled by user.")
    except Exception as e:
        print(f"\nAn unexpected error occurred: {str(e)}")
    finally:
        print("\nThank you for using the AI Report Generator!")


if __name__ == "__main__":
    main()